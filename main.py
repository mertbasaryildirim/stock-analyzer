import pandas as pd
import yfinance as yf
import mplfinance as mpf
from docx import Document
from docx.shared import Inches, RGBColor
import os
import time

# CSV dosyasını okuma
df = pd.read_csv('stocks.csv')
stock_symbols = df['stocks'].tolist()

# ./images klasörünü oluşturma
if not os.path.exists('./images'):
    os.makedirs('./images')

def calculate_ema(stock_data, period):
    return stock_data['Close'].ewm(span=period, adjust=False).mean()

def plot_stock_data(stock_symbol):
    stock_data = yf.Ticker(stock_symbol).history(period='1y')
    stock_data = stock_data.asfreq('B').ffill()  # Sadece iş günleri, tatil günlerini doldurma
    
    stock_data['EMA10'] = calculate_ema(stock_data, 10)
    stock_data['EMA25'] = calculate_ema(stock_data, 25)
    stock_data['EMA50'] = calculate_ema(stock_data, 50)
    stock_data['EMA100'] = calculate_ema(stock_data, 100)
    stock_data['EMA150'] = calculate_ema(stock_data, 150)
    
    # Renkleri tanımla
    colors = ['yellow', 'blue', 'green', 'purple', 'magenta']
    
    mpf.plot(stock_data, type='candle', style='charles', 
             mav=(10, 25, 50, 100, 150), 
             volume=False, 
             title=f'{stock_symbol} Price and EMAs',
             mavcolors=colors,
             savefig=f'./images/{stock_symbol}.png')

document = Document()
document.add_heading('Stock Analysis', 0)

summary = []

for stock_symbol in stock_symbols:
    try:
        plot_stock_data(stock_symbol)
        
        stock_data = yf.Ticker(stock_symbol).history(period='1y')
        stock_data = stock_data.asfreq('B').ffill()  # Sadece iş günleri, tatil günlerini doldurma
        ema10 = calculate_ema(stock_data, 10)
        ema25 = calculate_ema(stock_data, 25)
        ema50 = calculate_ema(stock_data, 50)
        ema100 = calculate_ema(stock_data, 100)
        ema150 = calculate_ema(stock_data, 150)
        
        latest_ema10 = ema10.iloc[-1]
        latest_ema25 = ema25.iloc[-1]
        latest_ema50 = ema50.iloc[-1]
        latest_ema100 = ema100.iloc[-1]
        latest_ema150 = ema150.iloc[-1]
        
        document.add_heading(stock_symbol, level=1)
        document.add_picture(f'./images/{stock_symbol}.png', width=Inches(4))
        
        if latest_ema10 > latest_ema25 > latest_ema50 > latest_ema100 > latest_ema150:
            color = RGBColor(0, 255, 0)  # Yeşil
            summary.append((stock_symbol, 'BUY'))
        elif latest_ema10 < latest_ema25 < latest_ema50 < latest_ema100 < latest_ema150:
            color = RGBColor(255, 0, 0)  # Kırmızı
            summary.append((stock_symbol, 'SELL'))
        else:
            color = RGBColor(255, 255, 0)  # Sarı
            summary.append((stock_symbol, 'HOLD'))
        
        p = document.add_paragraph()
        print(stock_symbol + " is calculated")
        time.sleep(5)
    
    except:
        print("Some kind of error with " + stock_symbol)

# Summary listesini BUY, SELL ve HOLD sırasına göre sırala
summary.sort(key=lambda x: {'BUY': 0, 'SELL': 1, 'HOLD': 2}[x[1]])

document.add_heading('Summary', level=1)
for stock, action in summary:
    action_color = {
        'BUY': RGBColor(0, 255, 0),
        'SELL': RGBColor(255, 0, 0),
        'HOLD': RGBColor(255, 255, 0)
    }
    action_rgb = action_color[action]
    
    p = document.add_paragraph()
    p.add_run(stock).font.color.rgb = RGBColor(0, 0, 0)  # Siyah renkte hisse sembolü
    p.add_run(f': {action}').font.color.rgb = action_rgb  # Renkli aksiyon

document.save('stock_analysis.docx')
